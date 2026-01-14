"""
从 .docx 文档自动生成 questions.js（选择题题库）
用法：
  python build_questions.py "填空题汇总附答案.docx" "questions.js"
"""
import json, re, random, sys
from pathlib import Path
from docx import Document

random.seed(42)

def clean_prefix(s: str) -> str:
    s = s.strip()
    s = re.sub(r"^\s*\d+\s*[\.、]\s*", "", s)
    s = re.sub(r"^\s*[（(]?\d+[）)]\s*", "", s)
    return s.strip()

def merge_paras(paras):
    def is_numbered(x): return bool(re.match(r"^\s*\d+\s*[\.、]", x))
    merged=[]
    for s in paras:
        if merged and (not is_numbered(s)) and (not re.search(r"[。！？]$", merged[-1])):
            merged[-1]=merged[-1]+" "+s
        else:
            merged.append(s)
    return merged

PATTERNS = [
    (re.compile(r"(常用端口号|端口号)\s*是\s*([^。；]+)"), 2),
    (re.compile(r"(长度分别是|分别是)\s*[:：]?\s*([^。；]+)"), 2),
    (re.compile(r"由\s*([^。；]+?)\s*组成"), 1),
    (re.compile(r"(主要包括|包括|包含)\s*([^。；]+)"), 2),
    (re.compile(r"通过\s*([^。；]+?)\s*来"), 1),
    (re.compile(r"需要\s*([^。；]+?)\s*(实现|进行|完成)"), 1),
    (re.compile(r"(可分为|分为|划分为)\s*([^。；]+)"), 2),
    (re.compile(r"(主要有)\s*([^。；]+)"), 2),
    (re.compile(r"有\s*[:：]\s*([^。；]+)"), 1),
    (re.compile(r"有\s*([^。；]+?)\s*(两种|三种|四种|五种|6种|六种)"), 1),
    (re.compile(r"最有影响的有\s*([^。；]+)"), 1),
    (re.compile(r"其中\s*([^，,。；]+?)\s*保持不变"), 1),
    (re.compile(r"([^，,。；]+?)\s*保持不变"), 1),
    (re.compile(r"被转换成\s*([^。；]+)"), 1),
    (re.compile(r"(即)\s*([^。；]+)"), 2),
    (re.compile(r"(表示为|表示是|表示)\s*([^。；]+)"), 2),
    (re.compile(r"(是)\s*([^。；]+)"), 2),
    (re.compile(r"(为)\s*([^。；]+)"), 2),
    (re.compile(r"(有)\s*([^。；]+)"), 2),  # 最后兜底
]

def extract_blank(clause: str):
    clause = clean_prefix(clause)
    clause = re.sub(r"\s+", " ", clause).strip()
    for pat, grp in PATTERNS:
        m = pat.search(clause)
        if m:
            ans = m.group(grp).strip().lstrip("：:")
            q = clause[:m.start(grp)] + "____" + clause[m.end(grp):]
            return q, ans, clause
    return None

def normalize_answer(a: str) -> str:
    a = a.strip().lstrip("：:").strip()
    a = re.sub(r"[。；;]$", "", a).strip()
    return a

NUM_RE = re.compile(r"^\d+(、\d+)*(B|b|字节|比特|位)?$")

def ans_type(a: str) -> str:
    a2 = a.strip()
    a_norm = re.sub(r"\s+", "", a2).replace("，", "、").replace(",", "、")
    if NUM_RE.fullmatch(a_norm) or re.fullmatch(r"\d+个?(字节|比特|位)", a_norm):
        return "num"
    if ("、" in a2) or ("和" in a2) or ("，" in a2) or ("," in a2) or (len(a2) >= 14):
        return "list"
    return "text"

def make_options(correct: str, pools: dict, all_answers: list[str]):
    t = ans_type(correct)
    pool = list({x for x in pools[t] if x != correct})
    random.shuffle(pool)
    distractors = pool[:3]
    if len(distractors) < 3:
        other = [x for tt in pools for x in pools[tt] if x != correct and x not in distractors]
        random.shuffle(other)
        distractors += other[:(3 - len(distractors))]
    opts = distractors + [correct]
    uniq = []
    for o in opts:
        if o not in uniq:
            uniq.append(o)
    opts = uniq
    while len(opts) < 4:
        cand = random.choice(all_answers)
        if cand != correct and cand not in opts:
            opts.append(cand)
    opts = opts[:4]
    random.shuffle(opts)
    return opts, opts.index(correct)

def build(docx_path: Path):
    doc = Document(str(docx_path))
    raw_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    paras = merge_paras(raw_paras)

    facts = []
    # 兼容协议表格：抽取端口 / 名称 / 功能
    for t in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        if not rows: 
            continue
        header = rows[0]
        for r in rows[1:]:
            if not r or not r[0]:
                continue
            proto = r[0].strip()
            name = r[1].strip() if len(r) > 1 else ""
            feature = r[2].strip() if len(r) > 2 else ""
            port = r[3].strip() if len(r) > 3 else ""
            if port:
                facts.append(f"{proto} 的常用端口号是 {port}。")
            if name:
                facts.append(f"{proto} 的中英文名称是 {name}。")
            if feature:
                facts.append(f"{proto} 的主要功能特征是：{feature}。")

    facts += paras

    clauses=[]
    for f in facts:
        for part in re.split(r"[；;]", f):
            part = part.strip()
            if part:
                clauses.append(part)

    qa=[]
    for c in clauses:
        res = extract_blank(c)
        if not res:
            continue
        q, a, orig = res
        qa.append({"question": q, "answer": normalize_answer(a), "explain": orig.strip()})

    all_answers = [x["answer"] for x in qa]
    pools = {"num": [], "list": [], "text": []}
    for a in all_answers:
        pools[ans_type(a)].append(a)

    questions=[]
    for i, item in enumerate(qa, start=1):
        opts, ci = make_options(item["answer"], pools, all_answers)
        questions.append({
            "id": i,
            "q": item["question"].strip(),
            "options": opts,
            "answerIndex": ci,
            "explain": item["explain"]
        })
    return questions

def main():
    if len(sys.argv) < 2:
        print("用法：python build_questions.py input.docx [output.js]")
        sys.exit(1)
    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2]) if len(sys.argv) >= 3 else Path("questions.js")
    questions = build(input_path)
    js = "/* 自动生成题库 */\nwindow.QUESTIONS = " + json.dumps(questions, ensure_ascii=False, indent=2) + ";\n"
    output_path.write_text(js, encoding="utf-8")
    print(f"已生成：{output_path}（{len(questions)}题）")

if __name__ == "__main__":
    main()
